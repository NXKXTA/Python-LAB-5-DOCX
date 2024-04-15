# для установки pip install python-docx
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# создание пустого документа
document = Document()

# добавление заголовка (по умолчанию level=1, 1-9)
document.add_heading('Дж. Р. Р. Толкин, level=1')
document.add_heading('Властилин колец, level=2', level=2)
document.add_heading('Часть I, level=9', level=9)
# заголовок со стилем для титульной страницы
document.add_heading('Братство кольца, level=0', level=0)

# добавляем параграфы
p1 = document.add_paragraph(
    'В Хоббитоне был переполох. Господин Бильбо Сумникс, хозяин Засумок, объявил о намерении '
    'отпраздновать свое стоодиннадцатилетие и пообещал очень щедрое угощение. '
    'Во всем Шире Бильбо слыл богатым чудаком с тех самых пор, как шестьдесят лет назад сначала '
    'запропал куда-то, а потом вернулся как снег на голову невесть откуда.')

p2 = document.add_paragraph(
    'О сокровищах, добытых Бильбо за тридевять земель, ходили неутихающие легенды. '
    'Многие верили, что подземелья Засумок ломятся от кладов. Но не только предполагаемое богатство '
    'заставляло хоббитов поглядывать на Бильбо с недоверчивым удивлением. '
    'Годы шли и шли, а по господину Сумниксу этого было не заметить. '
    'В девяносто он выглядел едва ли на пятьдесят. ')

document.add_paragraph('Абзац, с которым мы ничего не будем делать.')

# варианты форматирования
# выравнивание (LEFT, CENTER, RIGHT, JUSTIFY)
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# добавляем отступ для красной строки
p2.paragraph_format.first_line_indent = Mm(30)
# добавляем отступ слева
p2.paragraph_format.left_indent = Mm(50)
# добавляем отступ справа
p2.paragraph_format.right_indent = Mm(-30)
# расстояние между этим и предыдущим абзацем
p2.paragraph_format.space_before = Mm(20)
# расстояние между этим и следующим абзацем
p2.paragraph_format.space_after = Mm(40)
# межстрочный интервал
p2.paragraph_format.line_spacing = Mm(8)
#
# форматирование на основе прогона
p3 = document.add_paragraph(
    'В девяносто девять его называли “хорошо сохранившимся”, хотя правильнее было бы сказать ')
p3.add_run('“ничуть не изменившийся”').bold = True
p3.add_run('. Некоторые качали головами – дескать, многовато для одного,'
           'нечестно быть и очень богатым, и очень здоровым одновременно.')
p3_run = p3.add_run('“Это даром не пройдет, – говорили они, – вот увидите, как бы расплачиваться не пришлось”.')
p3_run.italic = True

# настраиваем шрифт
p3_run.font.name = 'Arial'
p3_run.font.size = Pt(24)
p3_run.font.color.rgb = RGBColor(0, 0, 255)

# добавление разрыва на страницу
document.add_page_break()

# добавляем и форматируем изображение
picture_paragraph = document.add_paragraph()
picture_run = picture_paragraph.add_run()
picture_run.add_picture('img.jpeg', width=Mm(100), height=Mm(70))
picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.save('LOTR.docx')
