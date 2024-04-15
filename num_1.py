from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# создание пустого документа
document = Document()

# Установка нового значения для левого поля (в дюймах)
section = document.sections[0]  # Получаем первый раздел документа
section.left_margin = Mm(15)

# Создание нового стиля для краткого описания гаража
style_for_description = document.styles['Normal']
style_for_description.font.name = 'Courier New'
style_for_description.font.size = Pt(14)
style_for_description.paragraph_format.first_line_indent = Mm(10)
style_for_description.paragraph_format.left_indent = Mm(20)
style_for_description.font.color.rgb = RGBColor(84, 32, 2)

# Создание нового стиля для параграфа списком того, что я предлагаю
style_for_p_list = document.styles['List Bullet 2']
style_for_p_list.font.name = 'Arial'
style_for_p_list.font.size = Pt(12)
style_for_p_list.italic = True
style_for_p_list.paragraph_format.left_indent = Mm(60)
style_for_p_list.font.color.rgb = RGBColor(242, 12, 39)

# Создание нового стиля для параграфа списком того, что покупатель отдаёт
style_for_p_coast = document.styles['List Bullet 3']
style_for_p_coast.font.name = 'Consolas'
style_for_p_coast.font.size = Pt(12)
style_for_p_coast.italic = True
style_for_p_coast.paragraph_format.left_indent = Mm(60)
style_for_p_coast.font.color.rgb = RGBColor(12, 242, 77)

"---------------------------------------------------------------------------------------------------------------"

# Добавление заголовка
head_1 = document.add_heading(level=1)
head_1_run = head_1.add_run('Продам гараж!!!')
head_1_run.bold = True
head_1_run.font.name = 'Courier New'
head_1_run.font.size = Pt(18)
head_1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Добавление параграфа 1
paragraph_1 = document.add_paragraph("Яма сухая. Внутри стеллажи металлические. Ворота в плохом состоянии."
                                     " Крыша не течёт."
                                     " Новая электропроводка и счётчик. Не кооперативный. В собственности.",
                                     style=style_for_description)

# Добавить параграф 2
paragraph_2 = document.add_paragraph("С меня:", style=style_for_description)

# Добавить параграф список
document.add_paragraph("Ржавый гараж", style=style_for_p_list)
document.add_paragraph("Куча бесполезного хлама", style=style_for_p_list)
document.add_paragraph("Крыша без дыр", style=style_for_p_list)
document.add_paragraph("")

# Добавить параграф 3
paragraph_3 = document.add_paragraph("С вас:", style=style_for_description)

# Добавить параграф список
document.add_paragraph("250 тысяч рублей", style=style_for_p_coast)

# добавляем и форматируем изображение
paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Путь к изображению
image_path = "./1.png"
# Добавление изображения в документ
run = paragraph.add_run()
inline_shape = run.add_picture(image_path, width=Mm(100), height=Mm(80))

# Сохраняем документ по указанному пути
document.save("C:\Фигня всякая\document.docx")
