import csv
import os
from docx import Document
from docxtpl import DocxTemplate

marathon_list = list()

marathon_path = './data_marathon.csv'
if not os.path.exists(marathon_path):
    print("Нет файла с марафоном")
    exit()

with open(marathon_path, 'r', encoding="utf-8") as csvfile:
    reader = csv.reader(csvfile, delimiter=',')
    for row in reader:
        marathon_data = list()
        marathon_data.append(int(row[0]))
        marathon_data = marathon_data + row[1:]
        marathon_list.append(marathon_data)

marathon_list.sort()

template = DocxTemplate("Shablon.docx")

document = Document()
last = marathon_list[0][0]
number = 0
for i in marathon_list:
    number += 1
    context = {
        'number': number,
        "year": i[0],
        "city": i[5],
        "name": ("-мужчины " if i[2] == "Male" else "-женщины ") + str(i[1]),
        "time": i[4]
    }
    if last != i[0]:
        document.add_page_break()

    template.render(context)
    temp_doc = template.docx
    for paragraph in temp_doc.paragraphs:
        document.add_paragraph(paragraph.text)

    last = i[0]

document.save("Marathon_results.docx")
